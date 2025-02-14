import streamlit as st
from PIL import Image
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, Mm
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from io import BytesIO

def handle_camera_tab(sections):
    """Xử lý chức năng camera và lưu trữ ảnh"""
    st.subheader("Báo cáo Camera")
    camera_images = {}
    
    for section in sections:
        st.write(f"### {section}")
        
        with st.container():
            if st.button(f"Thêm ảnh cho {section}", key=f"add_{section}"):
                if f"{section}_count" not in st.session_state:
                    st.session_state[f"{section}_count"] = 0
                st.session_state[f"{section}_count"] += 1
            
            if f"{section}_count" not in st.session_state:
                st.session_state[f"{section}_count"] = 1
            
            cols = st.columns(3)
            for i in range(st.session_state[f"{section}_count"]):
                with cols[i % 3]:
                    key = f"{section}_{i}"
                    camera_image = st.camera_input(
                        f"Ảnh {i+1} - {section}", 
                        key=f"camera_{key}"
                    )
                    
                    if camera_image:
                        try:
                            # Đọc ảnh dưới dạng bytes
                            image_bytes = camera_image.getvalue()
                            # Lưu trực tiếp bytes của ảnh
                            camera_images[key] = image_bytes
                            # Hiển thị preview
                            img = Image.open(BytesIO(image_bytes))
                            st.image(img, caption=f"Ảnh {i+1}", use_column_width=True)
                        except Exception as e:
                            st.error(f"Lỗi khi xử lý ảnh {key}: {str(e)}")
                            camera_images[key] = None
                    else:
                        if key not in camera_images:
                            camera_images[key] = None
            
            if st.session_state[f"{section}_count"] > 1:
                if st.button(f"Xóa ảnh cuối của {section}", key=f"remove_{section}"):
                    last_key = f"{section}_{st.session_state[f'{section}_count']-1}"
                    if last_key in camera_images:
                        del camera_images[last_key]
                    st.session_state[f"{section}_count"] -= 1
    
    return camera_images

def create_new_report(image_data):
    """Tạo báo cáo Word mới với layout linh hoạt"""
    document = Document()
    
    # Document setup
    section = document.sections[0]
    section.page_width = Inches(8.27)  # A4 width
    section.page_height = Inches(11.69)  # A4 height
    
    # Headers
    title = document.add_paragraph("GENERAL PHOTO")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True
    
    product_title = document.add_paragraph("TỔNG QUAN VỀ SẢN PHẨM/ PRODUCT OVERVIEW")
    product_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    product_title.runs[0].bold = True
    
    grape_title = document.add_paragraph("FRESH GREEN GRAPES")
    grape_title.runs[0].bold = True

    sections = [
        ("Tổng quan/Overview", "overview"),
        ("Kiểm tra khối lượng / Checking weight", "weight"),
        ("Kiểm tra kích cỡ / Checking size", "size"),
        ("Kiểm tra độ Brix / Checking Brix", "brix"),
        ("Kiểm tra độ cứng / Checking Firmness", "firmness"),
        ("Lỗi nghiêm trọng/ Serious defects", "serious"),
        ("Lỗi nặng/ Major defects", "major"),
        ("Lỗi nhẹ/ Minor defects", "minor"),
        ("Rụng cuống/ Shattering (Loosing) Berries", "shattering")
    ]

    for section_title, section_key in sections:
        # Add section header
        header = document.add_paragraph(section_title)
        header.runs[0].bold = True
        
        # Lấy tất cả ảnh của section này
        section_images = {k: v for k, v in image_data.items() if k.startswith(section_key) and v is not None}
        
        if section_images:
            # Tính số rows cần thiết (3 ảnh mỗi row)
            num_images = len(section_images)
            num_rows = (num_images + 2) // 3  # Ceiling division
            
            # Tạo table với số rows phù hợp
            table = document.add_table(rows=num_rows, cols=3)
            table.style = 'Table Grid'
            
            # Set column widths
            for col in table.columns:
                for cell in col.cells:
                    cell.width = Inches(2.5)
            
            # Add images
            for idx, (img_key, img_data) in enumerate(section_images.items()):
                try:
                    row = idx // 3
                    col = idx % 3
                    cell = table.cell(row, col)
                    paragraph = cell.paragraphs[0]
                    run = paragraph.add_run()
                    # Sử dụng trực tiếp bytes của ảnh
                    run.add_picture(BytesIO(img_data), width=Inches(2.3))
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception as e:
                    st.error(f"Lỗi khi chèn ảnh {img_key}: {str(e)}")
        
        # Add spacing after section
        document.add_paragraph()
        
        # Add specific text for defect sections
        if section_key == "serious":
            for _ in range(2):
                document.add_paragraph("Mô tả lỗi chung của nhóm - General").runs[0].bold = True
                document.add_paragraph("Thối - Rot")
        elif section_key == "major":
            document.add_paragraph("Lỗi nặng 3/ Major 3").runs[0].bold = True
            for _ in range(2):
                document.add_paragraph("Mô tả lỗi chung của nhóm - General").runs[0].bold = True
            for _ in range(12):
                document.add_paragraph("Dập nặng & mềm -- Major bruising & Soft")
            document.add_paragraph("Dập nặng -- Major bruising")
            document.add_paragraph("Mềm -- Soft")
        elif section_key == "minor":
            document.add_paragraph("Mô tả lỗi chung của nhóm - General").runs[0].bold = True
            document.add_paragraph("Sẹo-- Scars")
        elif section_key == "shattering":
            document.add_paragraph("Mô tả lỗi chung của nhóm - General").runs[0].bold = True
            for _ in range(2):
                document.add_paragraph("Rụng cuống / Shattering (Loosing) Berries")

    # Save document
    try:
        doc_stream = BytesIO()
        document.save(doc_stream)
        doc_stream.seek(0)
        return doc_stream
    except Exception as e:
        st.error(f"Lỗi khi lưu document: {str(e)}")
        return None

def render():
    st.header("CAMERA PHOTO")
    
    sections = [
        "Tổng Quan-Overview",
        "Checking Weight",
        "Checking Size",
        "Checking Brix",
        "Checking Firmness",
        "Serious Defects",
        "major 3",
        "minor Defects",
        "Shattering Berries"
    ]
    
    if "camera_images" not in st.session_state:
        st.session_state.camera_images = {}
    
    tab1, tab2 = st.tabs(["Camera Input", "Preview"])
    
    with tab1:
        camera_images = handle_camera_tab(sections)
        st.session_state.camera_images.update(camera_images)
        
        if st.button("Tạo báo cáo Word", type="primary"):
            with st.spinner("Đang tạo báo cáo..."):
                report_stream = create_new_report(st.session_state.camera_images)
                
                if report_stream:
                    st.success("Tạo báo cáo thành công!")
                    st.download_button(
                        label="Tải xuống báo cáo Word",
                        file_name="quality_inspection_report.docx",
                        data=report_stream,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                else:
                    st.error("Không thể tạo báo cáo. Vui lòng kiểm tra lại.")
    
    with tab2:
        st.subheader("Preview ảnh đã chụp")
        
        for section in sections:
            st.write(f"### {section}")
            section_images = {k: v for k, v in st.session_state.camera_images.items() 
                            if k.startswith(section) and v is not None}
            
            if section_images:
                cols = st.columns(3)
                for idx, (key, img_bytes) in enumerate(section_images.items()):
                    with cols[idx % 3]:
                        img = Image.open(BytesIO(img_bytes))
                        st.image(img, caption=f"Ảnh {idx+1}", use_column_width=True)

if __name__ == "__main__":
    render()
