import streamlit as st
from app.auth_manager import AuthManager
from app.sidebar_manager import SidebarManager
from pathlib import Path
import os
from openpyxl.utils import get_column_letter
import openpyxl.cell.cell
from app.absence_excel_processing import process_excel

# Cấu hình trang
st.set_page_config(
    page_title="Báo Cáo Vắng Mặt",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Kiểm tra quyền truy cập
auth_manager = AuthManager()
auth_manager.check_page_access("absence") # Cần chỉnh sửa trong auth_manager nếu cần

# Render sidebar
sidebar = SidebarManager()
sidebar.render_sidebar()

# Nếu chưa đăng nhập, chuyển hướng
if not st.session_state.get("authenticated", False):
    st.error("Trang này yêu cầu đăng nhập quản lý.")
    st.switch_page("pages/dashboard.py")

import pandas as pd
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Pt
import tempfile
import base64
from openpyxl import load_workbook
import logging

# Xử lý đường dẫn trong môi trường triển khai
if os.getenv('STREAMLIT_SERVER_PATH'):  # Môi trường Streamlit Cloud
    ROOT_DIR = Path('/mount/src/demoapp')  # Đường dẫn mặc định trên Streamlit Cloud
else:
    ROOT_DIR = Path(__file__).parent.parent.absolute()  # Môi trường local

# Cấu hình đường dẫn thư mục template
TEMPLATE_DIR = ROOT_DIR / "templates"

# Cấu hình đường dẫn các file template
TEMPLATE_FILES = {
    "Vắng Mặt Có Phép": TEMPLATE_DIR / "templates_1.docx",
    "Vắng Mặt Do Bệnh": TEMPLATE_DIR / "templates_2.docx",
    "Vắng Mặt Khác": TEMPLATE_DIR / "templates_3.docx", 
}

# Kiểm tra sự tồn tại của thư mục template
if not TEMPLATE_DIR.exists():
    st.error(f"Không tìm thấy thư mục template: {TEMPLATE_DIR}")
    TEMPLATE_DIR.mkdir(parents=True, exist_ok=True)

# Kiểm tra sự tồn tại của các file template
for attendance_type, template_path in TEMPLATE_FILES.items():
    if not template_path.exists():
        st.error(f"Không tìm thấy file template: {template_path}")

# Sửa đường dẫn file logo
LOGO_PATH = ROOT_DIR / "images" / "logo.png"

# Hàm chuyển đổi file ảnh thành Base64
def get_base64_image(image_path):
    try:
        with open(image_path, "rb") as img_file:
            encoded = base64.b64encode(img_file.read()).decode()
        return encoded
    except FileNotFoundError:
        st.error(f"Không tìm thấy file logo: {image_path}")
        return None

# Mã hóa Base64 cho logo
logo_base64 = get_base64_image(LOGO_PATH)

# Hiển thị logo nếu có
if logo_base64:
    # Căn giữa tiêu đề
    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        # Tiêu đề căn giữa
        st.markdown("""
            <h1 style='text-align: center;'>Báo Cáo Vắng Mặt</h1>
        """, unsafe_allow_html=True)

        # Căn giữa logo và tiêu đề phụ
        st.markdown(f"""
            <div style="display: flex; align-items: center; justify-content: center; margin-bottom: 20px;">
                <img src="data:image/png;base64,{logo_base64}" alt="Logo" style="margin-right: 10px; width: 30px; height: 30px;">
                <h3 style="margin: 0;">Eurofins SKHD</h3>
            </div>
        """, unsafe_allow_html=True)
else:
    # Căn giữa ngay cả khi không có logo
    st.markdown("""
        <h1 style='text-align: center;'>Báo Cáo Vắng Mặt</h1>
        <h3 style='text-align: center;'>Eurofins SKHD</h3>
    """, unsafe_allow_html=True)

st.markdown("---")

# Bước 1: Nhập thông tin
st.write("### Nhập Thông Tin")

# Khởi tạo các biến session
if 'confirmation_date' not in st.session_state:
    st.session_state['confirmation_date'] = datetime.now()
if 'grade' not in st.session_state:
    st.session_state['grade'] = "1"
if 'ma_hs' not in st.session_state:
    st.session_state['ma_hs'] = "1"
if 'hs_name' not in st.session_state:
    st.session_state['hs_name'] = ""

# Hiển thị khối lớp trên cùng một dòng
col1, col2 = st.columns(2)
with col1:
    grade = st.selectbox("Loại Hồ Sơ", ["Hồ Sơ 1", "Hồ Sơ 2", "Hồ Sơ 3"], key="grade_selectbox")[5:] # Lấy số từ chuỗi
with col2:
    ma_hs = st.text_input("Tên Hồ Sơ", key="ma_input")

# Hiển thị tên giáo viên và ngày xác nhận trên cùng một dòng
col3, col4 = st.columns(2)
with col3:
    hs_name = st.text_input("Người Phụ Trách", key="hs_input")
with col4:
    confirmation_date = st.date_input("Ngày Xác Nhận", st.session_state['confirmation_date'], key="confirmation_date_input")

# Nút chuyển sang bước tiếp theo
if st.button("Hoàn Tất Nhập", key="next_step_button"):
    st.session_state['confirmation_date_str'] = confirmation_date.strftime('%Y.%m.%d')
    st.session_state['grade'] = grade
    st.session_state['ma_hs'] = ma_hs
    st.session_state['hs_name'] = hs_name
    st.session_state['step'] = 2

# Bước 2: Tải lên và xử lý file Excel
if 'step' in st.session_state and st.session_state['step'] == 2:
    st.write("### Tải Lên File Excel")

    uploaded_file = st.file_uploader("Tải file Excel", type=["xlsx", "xls"], key="file_uploader")

    if uploaded_file is not None:
        try:
            # Kiểm tra tên file
            st.write("### Đang xử lý file:", uploaded_file.name)

            # Lưu file tải lên thành file tạm
            with tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx') as temp_file:
                temp_file.write(uploaded_file.getvalue())
                temp_path = temp_file.name

            # Sử dụng hàm process_excel mới
            data = process_excel(temp_path)

            # Xóa file tạm
            try:
                os.unlink(temp_path)
            except Exception as e:
                logging.warning(f"Lỗi khi xóa file tạm: {e}")

            # Thêm ngày xác nhận vắng mặt
            data['Ngày Xác Nhận'] = st.session_state.get('confirmation_date_str', '')

            # Kiểm tra và hiển thị dữ liệu
            if data.empty:
                st.error("Không có dữ liệu để xử lý.")
            else:
                st.session_state['processed_data'] = data
                data = st.data_editor(data, key='data_editor', use_container_width=True)
                st.session_state['processed_data'] = data

                col1, col2 = st.columns([1, 1], gap="small")
                with col1:
                    st.button("Quay Lại Bước 1", on_click=lambda: st.session_state.update({'step': 1}))
                with col2:
                    st.button("Tạo và Tải File DOCX", on_click=lambda: st.session_state.update({'step': 3}))

        except Exception as e:
            st.error(f"Lỗi trong quá trình xử lý file Excel: {e}")
            logging.error(f"Lỗi xử lý Excel: {str(e)}")

# Bước 3: Tạo và tải file DOCX
if 'step' in st.session_state and st.session_state['step'] == 3:
    st.write("### Tạo và Tải File DOCX")

    processed_data = st.session_state.get('processed_data', pd.DataFrame())

    if not processed_data.empty:
        st.write("Đang tạo file DOCX dựa trên dữ liệu đã xử lý.")

        for attendance_type, template_file_name in TEMPLATE_FILES.items():
            # Xử lý đặc biệt cho loại vắng mặt do bệnh
            if attendance_type == 'Vắng Mặt Do Bệnh':
                filtered_data = processed_data[processed_data['Lý Do'].str.contains('Bệnh.*Vắng', regex=True)] # Cần kiểm tra xem cột này có phải 'Lý Do' không
            else:
                filtered_data = processed_data[processed_data['Lý Do'] == attendance_type]  # Cần kiểm tra xem cột này có phải 'Lý Do' không

            if filtered_data.empty:
                continue

            template_path = os.path.join(TEMPLATE_DIR, template_file_name)

            if not os.path.exists(template_path):
                st.error(f"Không tìm thấy file template: {template_file_name}")
                continue

            output_docx = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')

            # Tải file template
            master_doc = Document(template_path)

            for idx, data_row in filtered_data.iterrows():
                # Áp dụng dữ liệu của từng học sinh vào template
                temp_doc = Document(template_path)
                replacements = {
                    "{1}": str(st.session_state['grade']),
                    "{2}": str(st.session_state['class_name']),
                    "{3}": str(int(data_row['Số Thứ Tự'])),
                    "{FullName}": str(data_row['Họ và Tên']),
                    "{Reason}": str(data_row['Lý Do']),
                    "{StartDate}": str(data_row['Ngày Bắt Đầu']),
                    "{Enđate}": str(data_row['Ngày Kết Thúc']),
                    "{NumberOfDate}": str(data_row['Số Ngày Vắng']),
                    "{ConfirmDate}": str(data_row['Ngày Xác Nhận']),
                    "{AdminName}": str(st.session_state['teacher_name'])
                }

                # Hàm thay thế text
                def replace_text_in_paragraph(paragraph, replacements):
                    for run in paragraph.runs:
                        text = run.text
                        for key, value in replacements.items():
                            if key in text:
                                text = text.replace(key, value)
                        run.text = text

                # Thay thế text trong văn bản
                for paragraph in temp_doc.paragraphs:
                    replace_text_in_paragraph(paragraph, replacements)

                # Thay thế text trong bảng
                for table in temp_doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                replace_text_in_paragraph(paragraph, replacements)

                # Thêm nội dung template đã sửa vào văn bản chính
                for element in temp_doc.element.body:
                    master_doc.element.body.append(element)

            # Xóa trang template đầu tiên
            if len(master_doc.element.body) > 1:
                master_doc.element.body.remove(master_doc.element.body[0])

            # Lưu file đã sửa
            master_doc.save(output_docx.name)

            # Cung cấp nút tải file trên Streamlit
            with open(output_docx.name, "rb") as f:
                st.download_button(
                    label=f"Tải {st.session_state['grade']} Khối {st.session_state['class_name']} Lớp {attendance_type} Báo Cáo Vắng Mặt({confirmation_date.month} Tháng).docx",
                    data=f,
                    file_name=f"{st.session_state['grade']} Khối {st.session_state['class_name']} Lớp {attendance_type} Báo Cáo Vắng Mặt({confirmation_date.month} Tháng).docx",
                    key=f"{attendance_type}_download"
                )
    else:
        st.error("Không có dữ liệu đã xử lý. Vui lòng kiểm tra lại dữ liệu ở bước trước.")

st.markdown("---")

st.subheader("Hướng Dẫn")
st.info(
    "Hệ thống Báo Cáo Vắng Mặt Thông Minh được thiết kế để đơn giản hóa công việc.\n\n"
    "Đường dẫn tải file Excel:\n"
    "[Naeis]-[Tình Hình và Thống Kê Điểm Danh]-[Tìm Kiếm]-[Tải Excel]\n\n"
    "Nếu có lỗi hoặc vấn đề, xin vui lòng thông báo!"
)

st.markdown("---")
st.markdown("<div style='text-align: right;'>Người thực hiện: Trung Vũ</div>", unsafe_allow_html=True)  